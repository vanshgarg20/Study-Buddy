# app.py - Study Buddy (final) 
# Features: Groq LLM remote generator + quizzes + resource links + onboarding + exports + chat assistant
# Requirements: pip install streamlit requests reportlab python-docx (optional)
import re, time, io, base64, hashlib, json, math
from typing import Dict, Any, List, Optional
from html import escape
import streamlit as st
import streamlit.components.v1 as components
import requests

# optional extras
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# --- small config ---
st.set_page_config(page_title="Study Buddy", layout="wide", initial_sidebar_state="expanded")
SHOW_DEBUG_DEFAULT = False     # hide banners/messages unless user enables debug

# ---------- session state ----------
if "demo_state" not in st.session_state:
    st.session_state["demo_state"] = {"plans": []}
if "view" not in st.session_state:
    st.session_state["view"] = "main"   # 'main' | 'onboarding' | 'chat'
if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []

def get_user_state(uid: str) -> Dict[str, Any]:
    return st.session_state["demo_state"]

def save_user_state(uid: str, state: Dict[str, Any]) -> None:
    st.session_state["demo_state"] = state

# ---------- helpers ----------
def slugify(s: Optional[str]) -> str:
    if s is None: return "item"
    s = str(s).strip().lower()
    s = re.sub(r'[^0-9a-z]+', '_', s)
    s = re.sub(r'__+', '_', s).strip('_')
    return s or "item"

def plan_to_text(record: Dict[str, Any]) -> str:
    lines = []
    lines.append(f"Study Buddy — Plan: {record.get('topic','')}")
    lines.append(f"Saved: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(record.get('created_at', time.time()))) }")
    lines.append(f"Source: {record.get('source','local')}")
    lines.append("")
    answers = record.get("answers", {})
    lines.append(f"Weeks: {record.get('weeks','—')}  •  Skill: {answers.get('skill_level','—')}  •  Hours/day: {answers.get('hours_per_day','—')}")
    if answers.get('goal'):
        lines.append(f"Goal: {answers.get('goal')}")
    lines.append("")
    plan = record.get("plan", {})
    weeks_list = plan.get("weeks", [])
    for w_i, w in enumerate(weeks_list, start=1):
        lines.append(f"Week {w_i}")
        for d in w.get("days", []):
            lines.append(f"  - {d.get('day')}: {', '.join(d.get('topics',[]))}")
        lines.append("")
    if plan.get("daily_template"):
        lines.append("Daily Template:")
        lines.append(f"  {plan.get('daily_template')}")
    if plan.get("resources"):
        lines.append("")
        lines.append("Resources:")
        for r in plan.get("resources", []):
            lines.append(f"  - {r.get('title')} — {r.get('url') if r.get('url') else ''}")
    if record.get('quiz'):
        lines.append("")
        lines.append("Quiz:")
        for i,q in enumerate(record.get('quiz',[]), start=1):
            lines.append(f"{i}. {q.get('q')}")
            for idx,opt in enumerate(q.get('options',[]), start=1):
                lines.append(f"   {idx}) {opt}")
            lines.append(f"   Answer: {q.get('a')}")
    return "\n".join(lines)

# ---------- local deterministic fallback generator (works without remote)
# lightweight domain detection to make plans relevant
def detect_domain_from_topic(topic: str) -> str:
    t = topic.lower()
    if any(k in t for k in ["interview","resume","soft","communication","behaviour","behavioral","hr","offer","mock"]):
        return "career"
    if any(k in t for k in ["data structure","dsa","algorithm","algorithms","leetcode","codeforce","algo"]):
        return "dsa"
    if any(k in t for k in ["machine learning","ml","deep learning","neural","model"]):
        return "ml"
    if any(k in t for k in ["web","react","frontend","backend","javascript","node"]):
        return "web"
    return "general"

def local_plan(topic: str, weeks: int, answers: Dict[str,Any]) -> Dict[str,Any]:
    domain = detect_domain_from_topic(topic)
    base = ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]
    weeks_list = []
    tokens = [w for w in re.split(r'\W+', topic.lower()) if w]
    seed = sum(ord(c) for c in topic) % 7 if topic else 3
    for w in range(max(1,weeks)):
        days = []
        for i,day in enumerate(base):
            token = tokens[(seed + i + w) % max(1,len(tokens))] if tokens else "practice"
            if domain == "career":
                topics = [f"Mock interview: {token}", "STAR answers", "Resume bullet refinement"]
            elif domain == "dsa":
                topics = [f"Topic concept: {token}", f"Practice problems: {token}", "Speed & pattern drills"]
            elif domain == "ml":
                topics = [f"Theory: {token}", f"Notebook experiment: {token}", "Read paper / blog"]
            else:
                topics = [f"Learn: {token}", "Practice exercises", "Mini-project task"]
            days.append({"day": day, "topics": topics})
        weeks_list.append({"days": days})
    daily_template = f"{answers.get('hours_per_day',2)} hours/day: mix of theory + practice"
    resources = [{"title": f"Search: {topic} on YouTube", "url": f"https://www.youtube.com/results?search_query={topic.replace(' ','+')}"}, {"title": f"Search: {topic} on Coursera", "url": f"https://www.coursera.org/search?query={topic.replace(' ','+')}" }]
    return {"weeks": weeks_list, "daily_template": daily_template, "resources": resources}

# ---------- remote Groq wrapper (LLM) ----------
def call_groq_api(model_name: str, system_prompt: str, user_prompt: str, max_tokens=1200, temp=0.2) -> Dict[str,Any]:
    api_key = st.secrets.get("API_KEY")
    api_url = st.secrets.get("API_URL")
    if not api_key or not api_url:
        raise RuntimeError("No API secrets")
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model_name, "messages":[{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}], "max_tokens":max_tokens, "temperature":temp}
    resp = requests.post(api_url, json=payload, headers=headers, timeout=20)
    resp.raise_for_status()
    return resp.json()

# generate remote plan (requests JSON-only output with quiz + resources)
def generate_plan_remote(uid: str, topic: str, weeks: int, answers: Dict[str,Any], debug=False) -> Dict[str,Any]:
    model = st.secrets.get("GROQ_MODEL") or "llama-3.1-8b-instant"
    sysp = (
        "You are a structured plan generator. ALWAYS respond with valid JSON ONLY (no markdown, no explanation). "
        "Output schema: {\"topic\":string, \"plan\": {\"weeks\": [ {\"days\":[{\"day\":\"Mon\",\"topics\":[...]}, ...] }, ...], \"daily_template\":string, \"resources\":[{\"title\":...,\"url\":...}, ...] }, \"quiz\": [ {\"q\":...,\"options\":[...],\"a\": \"correct option text\"}, ... ] }"
    )
    userp = (
        f"Create a concrete study plan for topic: \"{topic}\". Weeks: {weeks}. Skill: {answers.get('skill_level')}. "
        f"Hours per day: {answers.get('hours_per_day')}. Goal: {answers.get('goal')}. "
        "Make topics specific and actionable. Include a short quiz (3 questions) at the end with options and correct answer. "
        "Also include resource suggestions; for each resource include a reasonable url (youtube search link or official course link if known). Keep JSON compact."
    )
    try:
        data = call_groq_api(model, sysp, userp, max_tokens=1200, temp=0.2)
        # parse assistant text (OpenAI-like)
        choices = data.get("choices") if isinstance(data,dict) else None
        assistant_text = ""
        if choices and isinstance(choices, list) and choices:
            assistant_text = choices[0].get("message",{}).get("content","") or choices[0].get("text","")
        else:
            # maybe direct json
            assistant_text = json.dumps(data)
        # try parse JSON
        try:
            parsed = json.loads(assistant_text)
        except Exception:
            # extract first JSON-looking substring
            m = re.search(r'\{[\s\S]*\}', assistant_text)
            if m:
                parsed = json.loads(m.group(0))
            else:
                raise RuntimeError("Model did not return JSON")
        # basic validation
        if not parsed.get("plan"):
            raise RuntimeError("Remote plan missing 'plan' key")
        rec = {
            "topic": parsed.get("topic", topic),
            "weeks": int(max(1, weeks)),
            "answers": answers,
            "plan": parsed.get("plan"),
            "quiz": parsed.get("quiz", []),
            "created_at": int(time.time()),
            "source": f"remote ({model})"
        }
        return rec
    except Exception as e:
        if debug:
            st.error(f"Remote generation failed: {e}")
        raise

# wrapper with fallback & retries; will silently use local if not working unless debug True
def generate_plan(uid: str, topic: str, weeks: int, answers: Dict[str,Any], debug=False) -> Dict[str,Any]:
    # try remote if secrets set
    if st.secrets.get("API_KEY") and st.secrets.get("API_URL"):
        # attempt once with configured model
        model = st.secrets.get("GROQ_MODEL") or "llama-3.1-8b-instant"
        try:
            return generate_plan_remote(uid, topic, weeks, answers, debug=debug)
        except Exception as e:
            if debug:
                st.warning(f"Remote failed: {e} — falling back to local.")
    # fallback local
    return {"topic": topic, "weeks": int(max(1,weeks)), "answers": answers, "plan": local_plan(topic, weeks, answers), "quiz": [], "created_at": int(time.time()), "source": "local"}

# ---------- export helpers ----------
def make_md_bytes(record: Dict[str,Any]) -> bytes:
    txt = plan_to_text(record)
    return txt.encode("utf-8")

def make_ics_bytes(record: Dict[str,Any]) -> bytes:
    # create a simple recurring weekly schedule per plan: each week's Mon->Sun to separate events
    plan = record.get("plan",{})
    weeks = plan.get("weeks",[])
    start = time.gmtime(record.get("created_at", time.time()))
    dtstart = time.strftime("%Y%m%dT%H%M%SZ", start)
    # build events for each day in weeks, spaced by week offset
    ics_lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//StudyBuddy//EN"
    ]
    base_date = time.localtime(record.get("created_at", time.time()))
    # naive: assign each day to consecutive days starting today
    cur_ts = int(record.get("created_at", time.time()))
    for w_idx, wk in enumerate(weeks):
        for d in wk.get("days",[]):
            start_ts = cur_ts + (w_idx*7*86400)
            dt = time.gmtime(start_ts)
            dtstamp = time.strftime("%Y%m%dT%H%M%SZ", dt)
            uid = f"{record.get('created_at')}-{w_idx}-{d.get('day')}"
            ics_lines += [
                "BEGIN:VEVENT",
                f"UID:{uid}",
                f"DTSTAMP:{dtstamp}",
                f"DTSTART:{dtstamp}",
                f"SUMMARY:{record.get('topic')} — {d.get('day')}",
                f"DESCRIPTION:{', '.join(d.get('topics',[]))}",
                "END:VEVENT"
            ]
    ics_lines.append("END:VCALENDAR")
    return ("\n".join(ics_lines)).encode("utf-8")

def make_docx_bytes(record: Dict[str,Any]) -> Optional[bytes]:
    if not DOCX_AVAILABLE:
        return None
    doc = docx.Document()
    doc.add_heading(f"Study Buddy — {record.get('topic')}", level=1)
    doc.add_paragraph(plan_to_text(record))
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ---------- chat assistant (simple) ----------
def send_chat_message(prompt: str) -> str:
    model = st.secrets.get("GROQ_MODEL") or "llama-3.1-8b-instant"
    sysp = "You are a helpful assistant. Reply concisely."
    data = call_groq_api(model, sysp, prompt, max_tokens=400, temp=0.3)
    choices = data.get("choices", [])
    if choices and isinstance(choices[0], dict):
        return choices[0].get("message",{}).get("content","") or choices[0].get("text","")
    return str(data)

# ---------- UI (sidebar + controls) ----------
with st.sidebar:
    st.header("Create Plan")
    show_debug = st.checkbox("Show debug messages", value=SHOW_DEBUG_DEFAULT, help="Show remote debug banners and errors")
    view = st.radio("View:", ["Main", "Onboarding", "Chat"], index=0 if st.session_state["view"]=="main" else (1 if st.session_state["view"]=="onboarding" else 2))
    st.session_state["view"] = "main" if view=="Main" else ("onboarding" if view=="Onboarding" else "chat")
    st.markdown("---")

    # form inputs
    if "topic_input" not in st.session_state: st.session_state["topic_input"]="machine learning"
    if "weeks_input" not in st.session_state: st.session_state["weeks_input"]=2
    if "skill_input" not in st.session_state: st.session_state["skill_input"]="beginner"
    if "hours_input" not in st.session_state: st.session_state["hours_input"]=2
    if "goal_input" not in st.session_state: st.session_state["goal_input"]="Build a small project"
    if "username_input" not in st.session_state: st.session_state["username_input"]=""

    with st.form(key="create_form"):
        topic = st.text_area("Topic", key="topic_input", height=38)
        weeks = st.number_input("Weeks", min_value=1, max_value=52, value=st.session_state["weeks_input"], key="weeks_input")
        skill_level = st.selectbox("Skill level", ["beginner","intermediate","advanced"], index=["beginner","intermediate","advanced"].index(st.session_state["skill_input"]), key="skill_input")
        hours_per_day = st.slider("Hours per day", 1, 8, value=st.session_state["hours_input"], key="hours_input")
        goal = st.text_area("Goal (short)", key="goal_input", height=64)
        username = st.text_input("Username (optional)", key="username_input", help="Optional: same name keeps plans in session")
        submit = st.form_submit_button("Generate")

# top header / optionally small debug info (hidden unless show_debug)
st.markdown("<div style='padding:14px 20px;'></div>", unsafe_allow_html=True)
if show_debug:
    if st.secrets.get("API_KEY") and st.secrets.get("API_URL"):
        st.info("Remote generation: ENABLED (using API_KEY + API_URL from secrets).")
    else:
        st.info("Remote generation: NOT enabled — will use local generator.")

# onboarding page
if st.session_state["view"] == "onboarding":
    st.markdown("<div style='max-width:920px;margin:30px auto;'>", unsafe_allow_html=True)
    st.title("Welcome to Study Buddy")
    st.write("This app generates topic-aware study plans. Use the sidebar to enter a topic and click Generate.")
    st.write("Features included in this final version:")
    st.write("- LLM-powered plan generation (Groq) with local fallback")
    st.write("- Auto-generated short quiz (3 questions) per plan")
    st.write("- Resource suggestions with search links (YouTube/Coursera)")
    st.write("- Export to Markdown, DOCX (if available), and ICS calendar")
    st.write("- Simple chat assistant (Groq)")
    st.write("")
    st.write("Tips:")
    st.write(" - Put API_KEY and API_URL into Streamlit Secrets to enable remote generation.")
    st.write(" - If you prefer to keep debug messages hidden, uncheck 'Show debug messages' in the sidebar.")
    st.markdown("</div>", unsafe_allow_html=True)
    st.stop()

# chat view
if st.session_state["view"] == "chat":
    st.title("Assistant (Groq chat)")
    user_msg = st.text_area("Ask the assistant anything (short prompt)", height=80, key="chat_input")
    if st.button("Send"):
        if not user_msg.strip():
            st.warning("Please write a message.")
        else:
            try:
                resp = send_chat_message(user_msg)
                st.success("Assistant replied:")
                st.write(resp)
                st.session_state["chat_history"].append({"q": user_msg, "a": resp, "ts": int(time.time())})
            except Exception as e:
                st.error(f"Chat failed: {e}")
    if st.session_state["chat_history"]:
        st.markdown("### Chat history")
        for m in reversed(st.session_state["chat_history"]):
            st.markdown(f"**Q:** {m['q']}")
            st.markdown(f"**A:** {m['a']}")
            st.markdown("---")
    st.stop()

# normal main view
st.markdown("<div style='max-width:1100px;margin:10px auto;'>", unsafe_allow_html=True)
st.title("Study Buddy")
st.write("Topic-aware study plans — LLM used when available. (Debug banners are hidden by default.)")

# generate handling
uid = (st.session_state.get("username_input","").strip() or "session_user")
status_saved = False
if submit:
    if not st.session_state.get("topic_input","").strip():
        st.error("Please enter a topic.")
    else:
        answers = {"skill_level": st.session_state.get("skill_input"), "hours_per_day": st.session_state.get("hours_input"), "goal": st.session_state.get("goal_input")}
        # try remote first; debug mode passes show_debug
        try:
            rec = generate_plan(uid, st.session_state["topic_input"].strip(), int(st.session_state["weeks_input"]), answers, debug=show_debug)
        except Exception as e:
            if show_debug:
                st.error(f"Generation error: {e}")
            rec = {"topic": st.session_state["topic_input"].strip(), "weeks": int(st.session_state["weeks_input"]), "answers": answers, "plan": local_plan(st.session_state["topic_input"].strip(), int(st.session_state["weeks_input"]), answers), "quiz": [], "created_at": int(time.time()), "source": "local"}
        state = get_user_state(uid)
        plans = state.get("plans", [])
        # replace previous plan so page updates in-place
        if plans:
            plans[0] = rec
        else:
            plans.insert(0, rec)
        state["plans"] = plans
        save_user_state(uid, state)
        status_saved = True
        # do not show remote trial banners unless debug enabled (user requested)
        if show_debug:
            st.success(f"Plan created and saved ✅ (source: {rec.get('source')})")
        else:
            # normal UX: concise confirmation
            st.success("Plan ready ✅")

# render saved plans (only one top plan per user)
state = get_user_state(uid)
plans = state.get("plans", [])
if not plans:
    st.info("No saved plans yet. Create one from the sidebar.")
else:
    for idx,p in enumerate(plans):
        st.markdown("---")
        st.markdown(f"### {p.get('topic')}")
        st.markdown(f"_Saved: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(p.get('created_at',time.time())))} • source: {p.get('source')}_")
        # show meta
        answers = p.get("answers",{})
        st.markdown(f"- **Weeks:** {p.get('weeks')}  •  **Skill:** {answers.get('skill_level')}  •  **Hours/day:** {answers.get('hours_per_day')}")
        if answers.get("goal"):
            st.markdown(f"- **Goal:** {answers.get('goal')}")
        st.markdown("")
        # plan weeks
        weeks_list = p.get("plan",{}).get("weeks",[])
        for wi, wk in enumerate(weeks_list, start=1):
            st.markdown(f"**Week {wi}**")
            for d in wk.get("days",[]):
                st.markdown(f"- **{d.get('day')}:** {', '.join(d.get('topics',[]))}")
        # daily template & resources
        if p.get("plan",{}).get("daily_template"):
            st.markdown("**Daily Template:**")
            st.write(p.get("plan",{}).get("daily_template"))
        if p.get("plan",{}).get("resources"):
            st.markdown("**Resources:**")
            for r in p.get("plan",{}).get("resources"):
                title = r.get("title") if isinstance(r,dict) else str(r)
                url = r.get("url") if isinstance(r,dict) else None
                if url:
                    st.markdown(f"- [{title}]({url})")
                else:
                    st.markdown(f"- {title}")
        # quiz (if any)
        if p.get("quiz"):
            st.markdown("**Quiz:**")
            for i,q in enumerate(p.get("quiz",[]), start=1):
                st.markdown(f"{i}. {q.get('q')}")
                opts = q.get('options',[])
                for oi,opt in enumerate(opts, start=1):
                    st.markdown(f"   - ({oi}) {opt}")
            st.markdown("_Answers are hidden in UI; export to view solutions._")
        # actions: download md/docx/ics, copy plan
        col1,col2,col3,col4 = st.columns([0.28,0.28,0.28,0.16])
        with col1:
            md_bytes = make_md_bytes(p)
            b64 = base64.b64encode(md_bytes).decode()
            href = f"data:application/octet-stream;base64,{b64}"
            st.markdown(f"[Download .md]({href})")
        with col2:
            docx_bytes = make_docx_bytes(p)
            if docx_bytes:
                b64 = base64.b64encode(docx_bytes).decode()
                href = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"
                st.markdown(f"[Download .docx]({href})")
            else:
                st.markdown("DOCX not available (optional).")
        with col3:
            ics_bytes = make_ics_bytes(p)
            b64 = base64.b64encode(ics_bytes).decode()
            href = f"data:text/calendar;base64,{b64}"
            st.markdown(f"[Download .ics calendar]({href})")
        with col4:
            plan_text = plan_to_text(p)
            if st.button("Copy plan", key=f"copy_{idx}"):
                st.experimental_set_query_params()  # no-op to avoid warnings
                st.session_state["_clip"] = plan_text
                st.success("Plan copied to session — use clipboard method in browser.")
        st.markdown("---")

st.markdown("</div>", unsafe_allow_html=True)

# hide debug banners by default: show only if user enabled show_debug
if show_debug:
    st.info("Debug mode ON — remote errors and trial messages are visible above.")
else:
    # remove any leftover st.info/warning prints by not printing them; debugging only via sidebar toggle
    pass

# footer small
st.markdown("<div style='padding:8px 0 40px;'></div>", unsafe_allow_html=True)
